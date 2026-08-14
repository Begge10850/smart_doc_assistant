import os
from pathlib import Path

import pdfplumber
from docx import Document
import streamlit as st
import fitz  # PyMuPDF for annotation extraction
import boto3

from vision_engine import (
    NoReadableTextError,
    VisionProcessingError,
    extract_text_from_image_bytes,
)


SUPPORTED_FILE_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
}
IMAGE_FILE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
PDF_NATIVE_TEXT_MIN_CHARACTERS = 40
MAX_VISION_PDF_PAGES = 20


class DocumentProcessingError(RuntimeError):
    """A safe document-processing error that can be shown to an app user."""


# ─── AWS CREDENTIALS ────────────────────────────────────────────────────────────
try:
    aws_access_key = st.secrets["aws"]["AWS_ACCESS_KEY_ID"]
    aws_secret_key = st.secrets["aws"]["AWS_SECRET_ACCESS_KEY"]
except Exception:
    from dotenv import load_dotenv
    load_dotenv()
    aws_access_key = os.getenv("AWS_ACCESS_KEY_ID")
    aws_secret_key = os.getenv("AWS_SECRET_ACCESS_KEY")

s3 = boto3.client(
    's3',
    aws_access_key_id=aws_access_key,
    aws_secret_access_key=aws_secret_key,
    region_name="eu-north-1"
)

# ─── S3 FILE DOWNLOAD ───────────────────────────────────────────────────────────
def download_file_from_s3(file_name, download_path):
    try:
        s3.download_file("smart-doc-assistant-saidia", file_name, download_path)
        return True
    except Exception as e:
        print("Download error:", e)
        return False


# ─── LOCAL DOCUMENT INSPECTION ────────────────────────────────────────────────
def inspect_document(file_path):
    """Return safe document facts without sending the file to another service."""
    path = Path(file_path)
    extension = path.suffix.lower()

    if not path.is_file():
        raise DocumentProcessingError("The downloaded document could not be found.")

    if extension not in SUPPORTED_FILE_EXTENSIONS:
        raise DocumentProcessingError(
            f"The file type `{extension or 'unknown'}` is not supported."
        )

    if extension == ".pdf":
        document_kind = "pdf"
    elif extension == ".docx":
        document_kind = "word_document"
    elif extension == ".txt":
        document_kind = "text_document"
    else:
        document_kind = "image"

    metadata = {
        "file_name": path.name,
        "extension": extension,
        "document_kind": document_kind,
        "size_bytes": path.stat().st_size,
        "page_count": None,
        "image_width": None,
        "image_height": None,
    }

    if extension == ".pdf":
        try:
            with fitz.open(path) as document:
                metadata["page_count"] = document.page_count
        except Exception as exc:
            raise DocumentProcessingError(
                "The selected PDF is corrupted or could not be opened."
            ) from exc

    elif extension in IMAGE_FILE_EXTENSIONS:
        try:
            with fitz.open(path) as image_document:
                if image_document.page_count != 1:
                    raise DocumentProcessingError(
                        "The selected image has an unexpected page structure."
                    )

                image_page = image_document[0]
                image_pixmap = image_page.get_pixmap(alpha=False)
                metadata["page_count"] = 1
                metadata["image_width"] = image_pixmap.width
                metadata["image_height"] = image_pixmap.height
        except DocumentProcessingError:
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                "The selected image is corrupted or could not be opened."
            ) from exc

    return metadata


# ─── OPENAI VISION EXTRACTION ────────────────────────────────────────────────
def extract_image_with_vision(file_path):
    """Use OpenAI Vision to transcribe one uploaded image."""
    path = Path(file_path)
    mime_type = "image/png" if path.suffix.lower() == ".png" else "image/jpeg"

    return extract_text_from_image_bytes(
        path.read_bytes(),
        mime_type,
        page_label=path.name,
    )


def extract_scanned_pdf_with_vision(file_path):
    """Render and transcribe scanned-PDF pages with OpenAI Vision."""
    page_texts = []

    try:
        with fitz.open(file_path) as document:
            if document.page_count > MAX_VISION_PDF_PAGES:
                raise DocumentProcessingError(
                    f"Scanned PDFs are currently limited to "
                    f"{MAX_VISION_PDF_PAGES} pages to keep processing reliable."
                )

            for page_number, page in enumerate(document, start=1):
                # A 2x render improves small-text readability without giving the
                # model an unrestricted local path or direct access to S3.
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)

                try:
                    page_text = extract_text_from_image_bytes(
                        pixmap.tobytes("png"),
                        "image/png",
                        page_label=f"page {page_number}",
                    )
                except NoReadableTextError:
                    page_text = "[No readable text]"

                page_texts.append(f"[Page {page_number}]\n{page_text}")
    except (DocumentProcessingError, VisionProcessingError):
        raise
    except Exception as exc:
        raise DocumentProcessingError(
            "The scanned PDF could not be prepared for OpenAI vision."
        ) from exc

    readable_pages = [
        page_text for page_text in page_texts if "[No readable text]" not in page_text
    ]
    if not readable_pages:
        raise DocumentProcessingError(
            "OpenAI vision found no readable text in the scanned PDF."
        )

    return "\n\n".join(page_texts)


# ─── ANNOTATION EXTRACTION FROM PDF ─────────────────────────────────────────────
def extract_annotations_from_pdf(file_path):
    try:
        with fitz.open(file_path) as document:
            annotations = []

            for page in document:
                annotation = page.first_annot

                while annotation:
                    info = annotation.info
                    if info and info.get("content"):
                        annotations.append(info["content"])
                    annotation = annotation.next

        return "\n".join(annotations)

    except Exception as e:
        print("Annotation extraction failed:", e)
        return ""

# ─── TEXT EXTRACTION FROM PDF ───────────────────────────────────────────────────
def _extract_native_pdf_text(file_path):
    page_texts = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    page_texts.append(f"[Page {page.page_number}]\n{page_text}")
    except Exception as exc:
        print("pdfplumber extraction failed:", exc)

    return "\n\n".join(page_texts)


def _process_pdf(file_path):
    text = _extract_native_pdf_text(file_path)
    native_word_count = len(text.split())
    native_character_count = len("".join(text.split()))
    appears_scanned = native_character_count < PDF_NATIVE_TEXT_MIN_CHARACTERS

    if appears_scanned:
        text = extract_scanned_pdf_with_vision(file_path)
        extraction_method = "openai_vision"
    else:
        extraction_method = "native_pdf"

    try:
        annotations = extract_annotations_from_pdf(file_path)
        if annotations.strip():
            text += "\n\n[Annotations]\n" + annotations
    except Exception as exc:
        print("Annotation merge failed:", exc)

    return text, {
        "extraction_method": extraction_method,
        "used_vision": extraction_method == "openai_vision",
        "appears_scanned": appears_scanned,
        "native_word_count": native_word_count,
        "native_character_count": native_character_count,
    }


def extract_text_from_pdf(file_path):
    """Compatibility wrapper that returns only extracted PDF text."""
    text, _ = _process_pdf(file_path)
    return text


def _extract_docx_text(file_path):
    """Extract ordinary paragraphs and table rows from a Word document."""
    document = Document(file_path)
    text_parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cell_values = [cell.text.strip() for cell in row.cells]
            if any(cell_values):
                text_parts.append(" | ".join(cell_values))

    return "\n".join(text_parts)


# ─── MAIN HANDLER ───────────────────────────────────────────────────────────────
def process_document(file_path):
    """Inspect and extract a document, returning text plus decision metadata."""
    metadata = inspect_document(file_path)
    extension = metadata["extension"]

    try:
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as file_stream:
                text = file_stream.read()
            decision_metadata = {
                "extraction_method": "native_text",
                "used_vision": False,
                "appears_scanned": False,
            }

        elif extension == ".pdf":
            text, decision_metadata = _process_pdf(file_path)

        elif extension == ".docx":
            text = _extract_docx_text(file_path)
            decision_metadata = {
                "extraction_method": "native_docx",
                "used_vision": False,
                "appears_scanned": False,
            }

        elif extension in IMAGE_FILE_EXTENSIONS:
            text = extract_image_with_vision(file_path)
            decision_metadata = {
                "extraction_method": "openai_vision",
                "used_vision": True,
                "appears_scanned": True,
            }

        else:
            raise DocumentProcessingError("The selected file type is not supported.")
    except (DocumentProcessingError, VisionProcessingError):
        raise
    except UnicodeDecodeError as exc:
        raise DocumentProcessingError(
            "The selected text file is not valid UTF-8 text."
        ) from exc
    except Exception as exc:
        raise DocumentProcessingError(
            "The document could not be processed. Check the app logs."
        ) from exc

    text = text.strip()
    if not text:
        raise DocumentProcessingError(
            "No readable text could be extracted from the document."
        )

    metadata.update(decision_metadata)
    metadata["extracted_word_count"] = len(text.split())
    metadata["extracted_character_count"] = len(text)

    return {"text": text, "metadata": metadata}


def extract_text_from_file(file_path):
    """Compatibility wrapper that returns only extracted document text."""
    return process_document(file_path)["text"]
