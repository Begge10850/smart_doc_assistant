from docx import Document

from rag_pipeline import process_document


def test_docx_extracts_paragraphs_and_table_cells(tmp_path):
    """Regression test: DOCX extraction must include text stored in tables."""

    file_path = tmp_path / "incident_with_table.docx"

    document = Document()

    document.add_paragraph(
        "Incident INC-TEST-001 was reported by the customer."
    )

    table = document.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Tracking number"
    table.cell(0, 1).text = "NSP-DE-123456"

    table.cell(1, 0).text = "Declared value"
    table.cell(1, 1).text = "EUR 240.00"

    document.save(file_path)

    result = process_document(str(file_path))

    extracted_text = result["text"]

    # Ordinary paragraph must survive extraction.
    assert "INC-TEST-001" in extracted_text

    # Table content must also survive extraction.
    assert "Tracking number" in extracted_text
    assert "NSP-DE-123456" in extracted_text
    assert "Declared value" in extracted_text
    assert "EUR 240.00" in extracted_text

    # Verify that DOCX used native extraction rather than Vision.
    assert result["metadata"]["extraction_method"] == "native_docx"
    assert result["metadata"]["used_vision"] is False